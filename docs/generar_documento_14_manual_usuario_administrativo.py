import argparse
import re
import unicodedata
import zipfile
from datetime import date
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm


DOCS_DIR = Path(r"c:\Local\safi-ngac\docs")
MARKDOWN_PATH = DOCS_DIR / "documento_14_manual_usuario_administrativo.md"
TEMPLATE_PATH = DOCS_DIR / "doc_base.docx"
OUTPUT_DOCX_PATH = DOCS_DIR / "14-Manual-Usuario-Administrativo.docx"
IMAGES_DIR = DOCS_DIR / "images"
USER_IMAGE = IMAGES_DIR / "documento_14_manual_usuario_administrativo.png"
DOC_TITLE = "SAFI-NGAC"
DOC_SUBTITLE = "Manual de Usuario Administrativo"
PREFERRED_TABLE_STYLE = "Tabla de lista 4 - Énfasis 1"
PREFERRED_TABLE_STYLE_ALIASES = ["List Table 4 Accent 1"]


def ensure_dirs():
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)


def normalize_style_name(name):
    normalized = unicodedata.normalize("NFKD", name.strip().lower())
    normalized = "".join(char for char in normalized if not unicodedata.combining(char))
    return re.sub(r"\s+", " ", normalized)


def clean_heading_text(text):
    cleaned = re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", text).strip()
    cleaned = re.sub(r"^Documento\s+\d+\.?\s*", "", cleaned, flags=re.IGNORECASE).strip()
    return cleaned


def section_key(text):
    return normalize_style_name(clean_heading_text(text))


def clean_markdown_links(text):
    return re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)


class StyleCatalog:
    def __init__(self, document):
        self.by_name = {style.name: style for style in document.styles}
        self.by_normalized_name = {
            normalize_style_name(style.name): style.name for style in document.styles
        }
        self.table_styles = {
            style.name: style for style in document.styles if style.type == WD_STYLE_TYPE.TABLE
        }
        self.table_styles_by_normalized_name = {
            normalize_style_name(style.name): style.name
            for style in document.styles
            if style.type == WD_STYLE_TYPE.TABLE
        }
        self.character_styles = [
            style.name for style in document.styles if style.type == WD_STYLE_TYPE.CHARACTER
        ]

    def pick(self, *candidates, fallback="Normal"):
        for candidate in candidates:
            if not candidate:
                continue
            if candidate in self.by_name:
                return candidate
            normalized = self.by_normalized_name.get(normalize_style_name(candidate))
            if normalized:
                return normalized
        return fallback if fallback in self.by_name else "Normal"

    def pick_table(self, *candidates, fallback="Table Grid"):
        for candidate in candidates:
            if not candidate:
                continue
            if candidate in self.table_styles:
                return candidate
            normalized = self.table_styles_by_normalized_name.get(normalize_style_name(candidate))
            if normalized:
                return normalized
        return fallback if fallback in self.table_styles else next(iter(self.table_styles), None)


def get_preferred_table_style(styles):
    return styles.pick_table(PREFERRED_TABLE_STYLE, *PREFERRED_TABLE_STYLE_ALIASES, fallback="Table Grid")


def apply_inline_formatting(paragraph, text, styles):
    clean_text = clean_markdown_links(text)
    tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", clean_text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            style_name = styles.pick("Strong")
            if style_name in styles.character_styles:
                run.style = style_name
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            style_name = styles.pick("Emphasis")
            if style_name in styles.character_styles:
                run.style = style_name
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            style_name = styles.pick("Subtle Reference", "macro")
            if style_name in styles.character_styles:
                run.style = style_name
        else:
            paragraph.add_run(token.replace("`", ""))


def add_blank_line(document, styles):
    document.add_paragraph(style=styles.pick("Body Text", "Normal"))


def parse_markdown_tables(lines):
    blocks = []
    index = 0
    while index < len(lines):
        line = lines[index].rstrip("\n")
        if line.strip().startswith("|") and index + 1 < len(lines):
            separator = lines[index + 1].rstrip("\n")
            if separator.strip().startswith("|") and "---" in separator:
                header = [cell.strip() for cell in line.split("|")[1:-1]]
                rows = []
                index += 2
                while index < len(lines):
                    row_line = lines[index].rstrip("\n")
                    if not row_line.strip().startswith("|"):
                        break
                    rows.append([cell.strip() for cell in row_line.split("|")[1:-1]])
                    index += 1
                blocks.append(("table", header, rows))
                continue
        blocks.append(("line", line))
        index += 1
    return blocks


def extract_markdown_section(lines, target_key):
    in_section = False
    found = False
    collected = []
    for raw_line in lines:
        line = raw_line.rstrip("\n")
        header_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if header_match and len(header_match.group(1)) == 2:
            current_key = section_key(header_match.group(2))
            if in_section and current_key != target_key:
                break
            if current_key == target_key:
                in_section = True
                found = True
                continue
        if in_section:
            collected.append(raw_line)
    return collected if found else []


def replace_paragraph_text(paragraph, text):
    paragraph.clear()
    if text:
        paragraph.add_run(text)


def enable_update_fields(document):
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def set_table_width(table, total_width_cm=17.0):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total_width = Cm(total_width_cm)
    cols = len(table.columns) if table.columns else 1
    col_width = Cm(total_width_cm / cols)

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(total_width.twips)))

    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width


def generate_user_map(output_path):
    labels = ["Tablero", "Menú", "Gestión", "Contexto", "Validación"]
    values = [3, 5, 5, 4, 3]
    colors = ["#5B9BD5", "#2F5597", "#70AD47", "#ED7D31", "#A5A5A5"]

    fig, ax = plt.subplots(figsize=(12.0, 6.8))
    x_pos = np.arange(len(labels))
    ax.bar(x_pos, values, color=colors, width=0.58)
    ax.set_xticks(x_pos)
    ax.set_xticklabels(labels)
    ax.set_ylabel("Peso relativo en el uso")
    ax.set_title("SAFI-NGAC - Mapa base de navegación y uso administrativo")
    ax.grid(axis="y", linestyle="--", alpha=0.25)

    for index, value in enumerate(values):
        ax.text(index, value + 0.08, str(value), ha="center", fontsize=11, fontweight="bold")

    ax.text(
        -0.35,
        -0.95,
        "El uso funcional del sistema gira en torno al menú dinámico, el contexto y los módulos administrativos especializados.",
        fontsize=10,
    )
    fig.tight_layout()
    fig.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def update_cover_bindings(output_path, title, subject):
    temp_path = output_path.with_suffix(".tmp")
    title_xpath = "/ns1:coreProperties[1]/ns0:title[1]"
    subject_xpath = "/ns1:coreProperties[1]/ns0:subject[1]"

    def replace_bound_content(xml_text, xpath, value):
        escaped_xpath = re.escape(xpath)
        pattern = rf'(<w:dataBinding[^>]*w:xpath="{escaped_xpath}"[^>]*/>.*?<w:sdtContent>.*?<w:t(?: [^>]*)?>)(.*?)(</w:t>)'
        return re.sub(pattern, lambda match: match.group(1) + value + match.group(3), xml_text, count=1, flags=re.DOTALL)

    with zipfile.ZipFile(output_path, "r") as source_zip, zipfile.ZipFile(temp_path, "w") as target_zip:
        for info in source_zip.infolist():
            data = source_zip.read(info.filename)
            if info.filename == "word/document.xml":
                xml_text = data.decode("utf-8")
                xml_text = replace_bound_content(xml_text, title_xpath, title)
                xml_text = replace_bound_content(xml_text, subject_xpath, subject)
                data = xml_text.encode("utf-8")
            target_zip.writestr(info, data)

    temp_path.replace(output_path)


def update_info_table(document):
    if not document.tables:
        return

    table = document.tables[0]
    rows = table.rows
    if len(rows) < 5 or len(rows[0].cells) < 2:
        return

    rows[1].cells[0].text = "Nombre del Proyecto"
    rows[1].cells[1].text = DOC_TITLE

    rows[2].cells[0].text = "Nombre del Documento"
    rows[2].cells[1].text = DOC_SUBTITLE

    rows[3].cells[0].text = "Fecha del Documento"
    rows[3].cells[1].text = date.today().strftime("%d/%m/%Y")

    rows[4].cells[0].text = "Estado del Documento"
    rows[4].cells[1].text = "Aprobado para Documentación"

    table.style = get_preferred_table_style(StyleCatalog(document))
    set_table_width(table, 17.0)


def fill_intro_page(document, styles, markdown_path):
    lines = markdown_path.read_text(encoding="utf-8").splitlines(True)
    section_lines = extract_markdown_section(lines, "proposito y alcance")
    intro_texts = []
    for raw_line in section_lines:
        stripped = raw_line.strip()
        if not stripped or stripped == "---" or stripped.startswith("#") or stripped.startswith("|"):
            continue
        intro_texts.append(clean_markdown_links(stripped))

    intro_index = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip().lower() in {"introducción", "introduccion"}:
            intro_index = index
            break

    if intro_index is None:
        return

    blank_targets = []
    for position in range(intro_index + 1, len(document.paragraphs)):
        if document.paragraphs[position].text.strip():
            break
        blank_targets.append(position)

    for idx, text in enumerate(intro_texts[:3]):
        if idx < len(blank_targets):
            paragraph = document.paragraphs[blank_targets[idx]]
            replace_paragraph_text(paragraph, "")
            apply_inline_formatting(paragraph, text, styles)
        else:
            paragraph = document.add_paragraph(style=styles.pick("Body Text", "Normal"))
            apply_inline_formatting(paragraph, text, styles)

    document.add_picture(str(USER_IMAGE), width=Cm(16.0))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = document.add_paragraph(style=styles.pick("Caption", "Normal"))
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run("Mapa resumido de navegación y uso administrativo del sistema")
    document.add_page_break()


def get_default_table_style(document, styles):
    preferred = get_preferred_table_style(styles)
    if preferred:
        return preferred
    if document.tables and document.tables[0].style is not None:
        table_style = document.tables[0].style
        style_name = getattr(table_style, "name", table_style)
        resolved = styles.pick_table(style_name, fallback="")
        if resolved:
            return resolved
    return styles.pick_table("Table Grid")


def render_table(document, header, rows, default_style):
    add_blank_line(document, StyleCatalog(document))
    table = document.add_table(rows=1, cols=len(header))
    table.style = default_style
    set_table_width(table, 17.0)

    for index, value in enumerate(header):
        table.rows[0].cells[index].text = value

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            if index < len(cells):
                cells[index].text = value

    add_blank_line(document, StyleCatalog(document))


def should_skip_section(section_name):
    return section_name in {"control de cambios del documento", "proposito y alcance"}


def should_skip_heading(text):
    return section_key(text) == "objetivos del documento:"


def render_markdown(document, markdown_path, styles):
    blocks = parse_markdown_tables(markdown_path.read_text(encoding="utf-8").splitlines(True))
    skip_initial_title = True
    skip_initial_subtitle = True
    skip_section = False
    skip_heading_level = None
    in_conclusions = False
    default_table_style = get_default_table_style(document, styles)
    ordered_counters = [0, 0, 0]

    def reset_numbering():
        for index in range(len(ordered_counters)):
            ordered_counters[index] = 0

    for block in blocks:
        if block[0] == "table":
            if skip_section or skip_heading_level is not None or in_conclusions:
                continue
            _, header, rows = block
            render_table(document, header, rows, default_table_style)
            continue

        line = block[1].rstrip()
        if not line.strip() or line.strip() == "---":
            reset_numbering()
            continue

        header_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2).strip()
            current_key = section_key(text)

            if skip_heading_level is not None:
                if level <= skip_heading_level:
                    skip_heading_level = None
                else:
                    continue

            if level == 1 and skip_initial_title:
                skip_initial_title = False
                reset_numbering()
                continue
            if level == 2 and skip_initial_subtitle:
                skip_initial_subtitle = False
                reset_numbering()
                continue
            if should_skip_heading(text):
                skip_heading_level = level
                reset_numbering()
                continue
            if level == 2:
                if in_conclusions:
                    break
                skip_section = should_skip_section(current_key)
                if skip_section:
                    reset_numbering()
                    continue
                in_conclusions = current_key == "conclusiones"
            cleaned = clean_heading_text(text)
            heading_level = max(1, min(4, level - 1))
            reset_numbering()
            if heading_level == 1:
                document.add_page_break()
            else:
                add_blank_line(document, styles)
            paragraph = document.add_paragraph(style=styles.pick(f"Heading {heading_level}", f"Heading {min(level, 4)}"))
            apply_inline_formatting(paragraph, cleaned, styles)
            add_blank_line(document, styles)
            continue

        if skip_section or skip_heading_level is not None:
            reset_numbering()
            continue

        ordered_match = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if ordered_match:
            indent = min(len(ordered_match.group(1)) // 2, 2)
            ordered_counters[indent] += 1
            for index in range(indent + 1, len(ordered_counters)):
                ordered_counters[index] = 0
            paragraph = document.add_paragraph(style=styles.pick("List Paragraph", "Body Text", "Normal"))
            apply_inline_formatting(paragraph, f"{ordered_counters[indent]}. {ordered_match.group(3)}", styles)
            continue

        bullet_match = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        if bullet_match:
            reset_numbering()
            indent = min(len(bullet_match.group(1)) // 2, 2)
            paragraph = document.add_paragraph(style=styles.pick(f"List Bullet {indent + 1}", "List Bullet"))
            apply_inline_formatting(paragraph, bullet_match.group(2), styles)
            continue

        reset_numbering()
        paragraph = document.add_paragraph(style=styles.pick("Body Text", "Normal"))
        apply_inline_formatting(paragraph, line, styles)


def build_document(markdown_path, template_path, output_path):
    ensure_dirs()
    generate_user_map(USER_IMAGE)

    document = Document(str(template_path))
    styles = StyleCatalog(document)
    document.core_properties.title = DOC_TITLE
    document.core_properties.subject = DOC_SUBTITLE
    enable_update_fields(document)
    update_info_table(document)
    fill_intro_page(document, styles, markdown_path)
    render_markdown(document, markdown_path, styles)
    document.save(str(output_path))
    update_cover_bindings(output_path, DOC_TITLE, DOC_SUBTITLE)


def parse_args():
    parser = argparse.ArgumentParser(description="Genera el Documento 14 usando doc_base.docx como plantilla base.")
    parser.add_argument("--markdown", default=str(MARKDOWN_PATH), help="Ruta del Markdown fuente.")
    parser.add_argument("--template", default=str(TEMPLATE_PATH), help="Ruta de la plantilla doc_base.docx.")
    parser.add_argument("--output", default=str(OUTPUT_DOCX_PATH), help="Ruta de salida del DOCX.")
    return parser.parse_args()


def main():
    args = parse_args()
    markdown_path = Path(args.markdown)
    template_path = Path(args.template)
    output_path = Path(args.output)

    if not markdown_path.exists():
        raise FileNotFoundError(f"No existe el Markdown fuente: {markdown_path}")
    if not template_path.exists():
        raise FileNotFoundError(f"No existe la plantilla base: {template_path}")

    build_document(markdown_path, template_path, output_path)
    print(f"Documento generado en: {output_path}")
    print(f"Mapa generado en: {USER_IMAGE}")


if __name__ == "__main__":
    main()
